"""Generate a styled .docx acquisition prospectus for NeuralQuant.
Output: docs/SALE_PROSPECTUS.docx
Brand palette: deep navy + electric cyan accent + warm gold.

Tables-first, glance-readable. Single firm ask (Rs 3 Crore). Includes an
IP-safe System Architecture overview (sealed cores shown in gold).
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------- brand palette ----------
NAVY = RGBColor(0x0A, 0x1B, 0x2E)       # deep navy
CYAN = RGBColor(0x00, 0xB4, 0xD8)       # electric cyan accent
GOLD = RGBColor(0xC9, 0xA2, 0x27)       # warm gold
SLATE = RGBColor(0x3A, 0x4A, 0x5E)      # slate body
LIGHT = RGBColor(0xF2, 0xF5, 0xF8)      # light row
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK = RGBColor(0x12, 0x1A, 0x24)

NAVY_HEX = "0A1B2E"
CYAN_HEX = "00B4D8"
LIGHT_HEX = "EEF3F7"
GOLD_HEX = "C9A227"

CENTER = WD_ALIGN_PARAGRAPH.CENTER

# ---------- helpers ----------
def set_cell_shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, color="C7D2DC", sz="4"):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'single')
        e.set(qn('w:sz'), sz)
        e.set(qn('w:space'), '0')
        e.set(qn('w:color'), color)
        borders.append(e)
    tcPr.append(borders)

def style_run(run, size=11, bold=False, color=DARK, font="Calibri", italic=False):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    run.font.name = font

def shade_para(p, hex_fill):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_fill)
    pPr.append(shd)

def para_box_border(p, color=CYAN_HEX, sz="6"):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    for edge in ('top', 'left', 'bottom', 'right'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), sz)
        e.set(qn('w:space'), '6'); e.set(qn('w:color'), color)
        pbdr.append(e)
    pPr.append(pbdr)

def add_field(paragraph, field_code):
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve'); instrText.text = field_code
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar2)
    return run

def para(doc, text, size=11, bold=False, color=DARK, align=None, space_after=6, space_before=0, italic=False, font="Calibri"):
    p = doc.add_paragraph()
    if align is not None: p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    r = p.add_run(text); style_run(r, size, bold, color, font, italic)
    return p

def h1(doc, num, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(f"{num}.  {text}")
    style_run(r, 16, True, NAVY, "Calibri")
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '4'); bottom.set(qn('w:color'), CYAN_HEX)
    pbdr.append(bottom); pPr.append(pbdr)
    return p

def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    style_run(r, 12.5, True, SLATE, "Calibri")
    return p

def bullet(doc, text, bold_lead=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    if bold_lead:
        r1 = p.add_run(bold_lead); style_run(r1, 10.5, True, DARK)
        r2 = p.add_run(text); style_run(r2, 10.5, False, DARK)
    else:
        r = p.add_run(text); style_run(r, 10.5, False, DARK)
    return p

def checkbox(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.3)
    r0 = p.add_run("☐  "); style_run(r0, 11, False, CYAN, "Calibri")
    r = p.add_run(text); style_run(r, 10.5, False, DARK)
    return p

def styled_table(doc, headers, rows, col_widths=None, header_fill=NAVY_HEX, zebra=True):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for i, htext in enumerate(headers):
        cell = t.rows[0].cells[i]
        set_cell_shade(cell, header_fill)
        set_cell_borders(cell, color="0A1B2E", sz="4")
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = cell.paragraphs[0].add_run(htext)
        style_run(r, 10.5, True, WHITE, "Calibri")
        cell.paragraphs[0].paragraph_format.space_after = Pt(2)
        cell.paragraphs[0].paragraph_format.space_before = Pt(2)
    for ridx, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            c = cells[i]
            if zebra and ridx % 2 == 1:
                set_cell_shade(c, LIGHT_HEX)
            set_cell_borders(c)
            c.paragraphs[0].paragraph_format.space_after = Pt(2)
            c.paragraphs[0].paragraph_format.space_before = Pt(2)
            r = c.paragraphs[0].add_run(str(val))
            style_run(r, 10, False, DARK, "Calibri")
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    return t

def label_table(doc, rows, label_w=1.7, body_w=4.9, header_fill=NAVY_HEX):
    """Two-column label/value table where the left label cell is navy-shaded."""
    t = doc.add_table(rows=0, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for label, body in rows:
        cells = t.add_row().cells
        c0, c1 = cells[0], cells[1]
        set_cell_shade(c0, header_fill); set_cell_borders(c0, color="0A1B2E", sz="4")
        set_cell_borders(c1)
        c0.paragraphs[0].paragraph_format.space_after = Pt(2); c0.paragraphs[0].paragraph_format.space_before = Pt(2)
        c1.paragraphs[0].paragraph_format.space_after = Pt(2); c1.paragraphs[0].paragraph_format.space_before = Pt(2)
        r0 = c0.paragraphs[0].add_run(label); style_run(r0, 10, True, WHITE)
        r1 = c1.paragraphs[0].add_run(body); style_run(r1, 10, False, DARK)
    for row in t.rows:
        row.cells[0].width = Inches(label_w); row.cells[1].width = Inches(body_w)
    return t

def spacer(doc, pts=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(pts)
    p.paragraph_format.space_before = Pt(0)
    return p

# ---------- architecture diagram helpers ----------
def arch_box(doc, header, body=None, fill=NAVY_HEX, hcolor=CYAN, bcolor=WHITE, border=CYAN_HEX):
    p = doc.add_paragraph(); p.alignment = CENTER
    p.paragraph_format.space_after = Pt(2); p.paragraph_format.space_before = Pt(2)
    shade_para(p, fill); para_box_border(p, color=border, sz="6")
    rh = p.add_run(header); style_run(rh, 10.5, True, hcolor)
    if body:
        rh.add_break()
        rb = p.add_run(body); style_run(rb, 9, False, bcolor)
    return p

def arch_arrow(doc):
    p = doc.add_paragraph(); p.alignment = CENTER
    p.paragraph_format.space_after = Pt(0); p.paragraph_format.space_before = Pt(0)
    r = p.add_run("▼"); style_run(r, 11, True, SLATE)
    return p

def arch_engine_pair(doc):
    t = doc.add_table(rows=1, cols=2); t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.autofit = False
    data = [
        ("QUANT SCORING ENGINE",
         ["6-factor IRS% composite", "HMM regime-adaptive weighting"],
         "SEALED CORE — factor weights & IRS% formula disclosed only under NDA"),
        ("MULTI-AGENT AI COMMITTEE",
         ["PARA-DEBATE™:  bull · steelman-bear · synthesis", "reasoning LLM, function-calling tools"],
         "SEALED CORE — agent prompts & debate orchestration disclosed only under NDA"),
    ]
    for i, (title, lines, sealed) in enumerate(data):
        c = t.rows[0].cells[i]
        set_cell_shade(c, NAVY_HEX); set_cell_borders(c, color=CYAN_HEX, sz="6")
        p0 = c.paragraphs[0]; p0.alignment = CENTER
        p0.paragraph_format.space_after = Pt(2); p0.paragraph_format.space_before = Pt(2)
        r = p0.add_run(title); style_run(r, 10.5, True, CYAN)
        for ln in lines:
            pl = c.add_paragraph(); pl.alignment = CENTER; pl.paragraph_format.space_after = Pt(1)
            rr = pl.add_run(ln); style_run(rr, 9, False, WHITE)
        ps = c.add_paragraph(); ps.alignment = CENTER
        ps.paragraph_format.space_before = Pt(3); ps.paragraph_format.space_after = Pt(2)
        shade_para(ps, GOLD_HEX)
        rs = ps.add_run(sealed); style_run(rs, 8.5, True, DARK)
    for cell in t.rows[0].cells:
        cell.width = Inches(3.25)
    return t

# ---------- build document ----------
doc = Document()

for s in doc.sections:
    s.top_margin = Inches(0.8)
    s.bottom_margin = Inches(0.8)
    s.left_margin = Inches(0.9)
    s.right_margin = Inches(0.9)

normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)
normal.font.color.rgb = DARK

# ===== COVER PAGE =====
band = doc.add_paragraph()
band.paragraph_format.space_after = Pt(0)
pPr = band._p.get_or_add_pPr()
shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), NAVY_HEX); pPr.append(shd)
rb = band.add_run("  "); style_run(rb, 6, False, NAVY)

para(doc, "CONFIDENTIAL  ·  ACQUISITION PROSPECTUS", size=10, bold=True, color=CYAN,
     align=CENTER, space_before=120, space_after=4)

title = doc.add_paragraph(); title.alignment = CENTER
title.paragraph_format.space_after = Pt(2)
r = title.add_run("NeuralQuant"); style_run(r, 44, True, NAVY, "Calibri")

sub = doc.add_paragraph(); sub.alignment = CENTER
sub.paragraph_format.space_after = Pt(2)
r = sub.add_run("AI Stock-Intelligence Platform"); style_run(r, 18, False, SLATE, "Calibri")

rule = doc.add_paragraph(); rule.alignment = CENTER
rule.paragraph_format.space_after = Pt(14); rule.paragraph_format.space_before = Pt(8)
pPr = rule._p.get_or_add_pPr()
pbdr = OxmlElement('w:pBdr')
bottom = OxmlElement('w:bottom')
bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '18')
bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), GOLD_HEX)
pbdr.append(bottom); pPr.append(pbdr)
rr = rule.add_run("    "); style_run(rr, 2, False, WHITE)

para(doc, "Live  ·  Production-Deployed  ·  Investor-Demo-Ready",
     size=11, bold=False, color=SLATE, align=CENTER, space_after=180)

meta = doc.add_table(rows=4, cols=2)
meta.alignment = WD_TABLE_ALIGNMENT.CENTER
meta_data = [
    ("Prepared", "June 2026"),
    ("Seller", "Satyam Das  —  satyamdas03@gmail.com"),
    ("Asset", "NeuralQuant  (neuralquant.co)"),
    ("Status", "Live, production-deployed, v4.1.0"),
]
for i, (k, v) in enumerate(meta_data):
    c0 = meta.rows[i].cells[0]; c1 = meta.rows[i].cells[1]
    c0.width = Inches(1.4); c1.width = Inches(4.2)
    set_cell_borders(c0, color="FFFFFF", sz="2"); set_cell_borders(c1, color="FFFFFF", sz="2")
    r0 = c0.paragraphs[0].add_run(k.upper()); style_run(r0, 9, True, CYAN)
    r1 = c1.paragraphs[0].add_run(v); style_run(r1, 10.5, False, DARK)
    c0.paragraphs[0].paragraph_format.space_after = Pt(4)
    c1.paragraphs[0].paragraph_format.space_after = Pt(4)

para(doc, "Shareable with serious prospective acquirers under NDA.",
     size=9, italic=True, color=SLATE, align=CENTER, space_before=180, space_after=0)

doc.add_page_break()

# ===== BODY FOOTER (page numbers) =====
body_section = doc.sections[0]
footer = body_section.footer
fp = footer.paragraphs[0]
fp.alignment = CENTER
r = fp.add_run("NeuralQuant — Acquisition Prospectus   |   Confidential   |   Page ")
style_run(r, 8, False, SLATE)
add_field(fp, "PAGE")
r2 = fp.add_run(" of "); style_run(r2, 8, False, SLATE)
add_field(fp, "NUMPAGES")

# ===== 1. EXECUTIVE SUMMARY =====
h1(doc, 1, "Executive Summary")
para(doc, "NeuralQuant delivers what no competitor currently offers at a retail price point: institutional-grade quantitative factor analysis combined with an adversarial multi-agent AI research committee and voice-native interfaces — covering both Indian (NSE/BSE) and US (NYSE/NASDAQ) equity markets in a single production-deployed platform.",
     size=11, space_after=6)
para(doc, "Proprietary internals — source code, agent prompts, factor weights, and credentials — are disclosed only under NDA during due diligence. This document conveys scope and value, not the recipe.",
     size=10.5, color=SLATE, space_after=8)
label_table(doc, [
    ("What you are buying", "A complete, running business-in-a-box: domain, brand, codebase, 7-service infrastructure, data pipeline, pending patent, PARA-DEBATE™ methodology, backtest track record, and operations documentation. Live. Generating outputs. Ready for Day-1 operation by the acquirer."),
    ("Why it exists", "Built by a solo founder (Master of AI, UTS Sydney) to solve a real gap: no platform gave retail investors institutional-grade multi-factor quantitative research — and none covered India + US in a single AI-native framework."),
    ("Why it is being sold", "The seller is mid-programme in a Master of AI at the University of Technology Sydney. The platform deserves a full-time commercial team, enterprise sales, and institutional backing to reach its potential. It is being sold to a buyer who can scale it — not wound down."),
    ("Why now", "A structured sale is underway; multiple parties have been approached; the seller reserves the right to accept an LOI at any time. neuralquant.dev — a US competitor on the same brand — is actively building in this space. Early acquisition locks in the India + US dual-market moat before the category becomes contested."),
])

# ===== 2. THE COMPETITIVE MOAT =====
h1(doc, 2, "The Competitive Moat — Why NeuralQuant Wins")
para(doc, "No existing retail-accessible product combines all four of NeuralQuant's pillars simultaneously: (1) multi-factor quantitative scoring, (2) an adversarial multi-agent AI committee, (3) India + US dual-market native coverage, and (4) voice-native interfaces.", space_after=8)
styled_table(doc,
    ["Platform", "Price / mo", "Multi-agent AI", "Adversarial debate", "India (NSE-native)", "Voice agents", "Live-data quant scoring"],
    [
        ["Bloomberg Terminal", "$2,000", "✗", "✗", "✗", "✗", "Data only"],
        ["Danelfin", "~$34", "Scoring only", "✗", "✗ (US only)", "✗", "✓ (US only)"],
        ["Ticker.in", "~$5", "✗", "✗", "✓ (India only)", "✗", "Basic"],
        ["ChatGPT / Claude", "$20", "General", "✗", "✗", "General voice", "✗ (no live data)"],
        ["NeuralQuant", "$9.99", "✓", "✓", "✓", "✓ (2 live)", "✓ (US + India)"],
    ],
    col_widths=[1.3, 0.7, 0.95, 0.95, 0.95, 0.85, 1.05])
spacer(doc, 6)
para(doc, "Bloomberg has the data but no AI reasoning layer and a 200x higher price. Danelfin has scoring but is US-only, single-score, no adversarial scrutiny, no voice. Ticker.in covers India but has no AI engine. General-purpose LLMs have no live market data and no quant engine — they hallucinate numbers. NeuralQuant is the only product that combines all four pillars, at a retail price point.", size=10.5, space_after=6)
para(doc, "And building it yourself? The Engineering Investment section quantifies it: 24–37 senior engineer-months plus 6–12 months of calendar, the data-source and voice-stack integration scars a live system has already absorbed, and a pending patent that raises a copier's cost and legal risk. Acquisition collapses that timeline to zero.", size=10.5, color=SLATE, space_after=6)

# ===== 3. WHAT IS BEING ACQUIRED =====
h1(doc, 3, "What Is Being Acquired")
para(doc, "A complete, running business-in-a-box: domain, brand, codebase, infrastructure, data pipeline, intellectual property, and operating services.", space_after=8)
h2(doc, "Surface area at a glance")
styled_table(doc,
    ["Dimension", "Scale"],
    [
        ["Backend API routers", "33"],
        ["Web pages / routes", "38"],
        ["Running services", "7  (4 Render, 1 Vercel, 1 Railway, 1 Supabase)"],
        ["Scheduled jobs", "4 daily cron jobs"],
        ["Production tests", "116+ (passing)"],
        ["Markets covered", "US + India"],
        ["Stock universe", "~949 stocks  (502 India + 447 US)"],
        ["AI agents", "8 distinct agents (research, debate, voice)"],
        ["Live voice interfaces", "2  (voice PM + ambient companion)"],
        ["Patent", "Application filed (provisional stage)"],
        ["Current version", "v4.1.0"],
        ["Day 1 for the acquirer", "Platform runs autonomously — 4 cron jobs execute, 13/13 smoke suite passes, all 7 services live. Take ownership and it keeps running with zero immediate intervention."],
    ],
    col_widths=[1.9, 4.7])
spacer(doc, 8)

# ===== 4. FLAGSHIP CAPABILITIES =====
h1(doc, 4, "Flagship Capabilities — All Live")
para(doc, "Every capability below is live in production at neuralquant.co today — demonstrable in a 45-minute live walkthrough. ★ marks the three patent-pending IP differentiators.", space_after=8)
styled_table(doc,
    ["Capability", "What it does"],
    [
        ["Adversarial research committee  ★",
         "Multi-agent LLM investment committee producing a bull case, a steelmanned adversarial bear case (Devil's Advocate agent), and a synthesised verdict per stock. Designed to prevent consensus herding — the structural failure of single-model AI analysis. No competitor currently ships an adversarial debate architecture."],
        ["IRS% quantitative scoring engine  ★",
         "Proprietary Investment-Readiness Score (IRS%) combining 5 academic factors: Piotroski F-Score (quality), Jegadeesh-Titman momentum, HMM regime-adaptive weighting, a value composite, and a low-volatility factor. India equities add NSE Bhavcopy delivery_pct as a sixth liquidity-conviction signal — unique to Indian market microstructure, not offered by any comparable platform."],
        ["HMM regime detection  ★",
         "Hidden Markov Model classifying market state into 4 regimes (Risk-On / Risk-Off / Bear / Late-Cycle). Factor weights dynamically reweight by detected regime — Bear weights quality and value; Risk-On weights momentum. No retail AI research platform offers regime-adaptive quantitative scoring."],
        ["Voice portfolio manager",
         "Real-time conversational voice agent (LiveKit WebRTC + Deepgram STT + reasoning LLM + ElevenLabs TTS). ~20 function-calling tools, file upload, and a whiteboard. Full portfolio-discussion capability."],
        ["Ambient voice companion",
         "Always-listening voice assistant with page-context awareness, a wake word, and a morning market briefing. Voice-native equity research is ahead of the market — competitors are text-only."],
        ["Written research analyst",
         "Queryable AI analyst returning cited, data-validated written research on demand. Live price injection, clarification questions, and numeric-hallucination suppression via per-agent reconciliation against a [VERIFIED] data layer."],
        ["Live paper-trading dashboard",
         "Real-time trading matrix with equity curve, trade tape, and a strategy-reflection SSE stream. Dry-run operational; ready for live-trading activation."],
        ["Methodology & backtest page",
         "Public methodology page with legally-cautious disclosure. Backtest baseline stored in Supabase for full reproducibility — 87–91% hit rate, +12.69 to +14.83% alpha vs NIFTY50 (Q1 FY27)."],
    ],
    col_widths=[2.0, 4.6])

# ===== 5. SYSTEM ARCHITECTURE =====
h1(doc, 5, "System Architecture (Overview)")
para(doc, "How the platform fits together, end to end — without exposing the proprietary core. Data flows top to bottom. The two engines shown in gold are sealed: their internals (factor weights, HMM regime-transition matrix, agent prompts, and the IRS% composite formula) are disclosed only under NDA during due diligence. The quant engine and the AI committee cross-feed: quantitative scores frame the agents' debate, and the committee's verdict is reconciled against the [VERIFIED] data layer.", space_after=10)

arch_box(doc, "DATA SOURCES  (6+)", "FMP · Finnhub · yfinance · OpenBB · EDGAR · FRED   +   NSE Bhavcopy (India microstructure)")
arch_arrow(doc)
arch_box(doc, "INGESTION & NORMALIZATION", "US + India parity  ·  caching  ·  4 nightly cron jobs  ·  [VERIFIED] data layer")
arch_arrow(doc)
arch_engine_pair(doc)
arch_arrow(doc)
arch_box(doc, "CORE API", "FastAPI  ·  33 routers  ·  auth  ·  quota  ·  rate-limit  ·  sessions")
arch_arrow(doc)
arch_box(doc, "DELIVERY LAYER",
         "Web app (Next.js, 38 pages, PWA)   ·   Voice layer (LiveKit · STT · TTS · 2 agents)   ·   Live paper-trading engine (Railway · SSE)",
         fill=LIGHT_HEX, hcolor=NAVY, bcolor=SLATE, border="0A1B2E")
spacer(doc, 6)
para(doc, "What is deliberately not shown: per-regime factor weights, HMM transition probabilities, the IRS% composite formula, and the PARA-DEBATE™ agent prompts. These are the asset's trade secrets — conveyed only under NDA. The diagram lets a buyer understand how value is produced without enabling replication.",
     size=9.5, italic=True, color=SLATE, space_after=6)

# ===== 6. STRATEGIC FIT BY BUYER TYPE =====
h1(doc, 6, "Strategic Fit by Buyer Type")
para(doc, "NeuralQuant is not a one-size-fits-all asset — its value is highest to four distinct buyer profiles, each acquiring a different strategic advantage.", space_after=8)
styled_table(doc,
    ["Buyer type", "What they acquire", "Why it matters"],
    [
        ["Indian brokers", "Ready-made AI research layer for 90M+ Demat holders — no build required", "Differentiated retention feature deployable on day one vs an 18-month internal build"],
        ["US / global fintechs", "India market entry (NSE + BSE with India-specific signals)", "The dual-market pipeline would take ~18 months to build independently; here it is already live"],
        ["PE firms / family offices", "Proven alpha (+12.69% to +14.83% vs NIFTY50), patent-protected", "A defensible, IP-encumbered asset to acquire and scale with a commercial team"],
        ["AI companies", "Voice-native equity research — 2 live voice agents already deployed", "Voice-first financial research is ahead of the market; competitors are text-only"],
    ],
    col_widths=[1.6, 2.7, 2.3])
spacer(doc, 8)

# ===== 7. TECHNOLOGY STACK =====
h1(doc, 7, "Technology Stack")
styled_table(doc,
    ["Layer", "Technology"],
    [
        ["Backend", "FastAPI, Python 3.12, async"],
        ["Frontend", "Next.js 16, React 19, Tailwind v4 (PWA)"],
        ["Database / Auth", "Supabase  (Postgres + cookie-session auth)"],
        ["Voice real-time", "LiveKit Cloud  (WebRTC SFU)"],
        ["Speech-to-text", "Deepgram"],
        ["Text-to-speech", "ElevenLabs  (neural voices)"],
        ["LLM providers", "Anthropic (Claude) + AWS Bedrock (cross-region)"],
        ["Market data", "FMP Premium, Finnhub, yfinance, OpenBB, EDGAR, FRED"],
        ["Hosting", "Render (4), Vercel (web), Railway (trading), Supabase (DB)"],
        ["CI/CD", "GitHub Actions, Render + Vercel auto-deploy"],
        ["Repository", "Monorepo, uv workspace"],
    ],
    col_widths=[1.8, 4.8])

# ===== 8. INFRASTRUCTURE =====
h1(doc, 8, "Infrastructure (Operating Today)")
para(doc, "All services are running and verified as of the date of this document.", space_after=6)
styled_table(doc,
    ["Service", "Role", "State"],
    [
        ["Core API", "33-router FastAPI backend", "Live"],
        ["Voice agent worker", "LiveKit voice agents (2)", "Live"],
        ["Paper-trading worker", "Trading daemon", "Live (dry-run)"],
        ["Market-data proxy", "OpenBB platform (live-price unlock)", "Live"],
        ["Trading engine", "Crypto paper trading + SSE stream", "Live"],
        ["Web app", "neuralquant.co, 38 pages", "Live"],
        ["Database", "Postgres + auth", "Live"],
    ],
    col_widths=[1.8, 3.4, 1.4])
spacer(doc, 6)
para(doc, "Scheduled jobs: 4 daily cron jobs handling nightly scoring, market refresh, and per-market end-of-day wrap reports.", size=10.5, space_after=6)
para(doc, "Security hardening: the platform has undergone a multi-phase security pass including log redaction, secret-scanning in CI, row-level security policies, IDOR remediation, content-security-policy headers, rate-limiting fuses, upload guards, dependency auditing, and an audit-event log with an incident-response runbook.", size=10.5, color=SLATE, space_after=6)

# ===== 9. VALIDATION & TRACK RECORD =====
h1(doc, 9, "Validation & Track Record")
h2(doc, "Backtest results (Q1 FY27, India)")
styled_table(doc,
    ["Metric", "Result"],
    [
        ["Benchmark", "NIFTY50  (−6.38% over the period)"],
        ["Model pools vs benchmark", "All three pools beat NIFTY50"],
        ["Alpha (outperformance)", "+12.69% to +14.83%"],
        ["Hit rate", "87–91%"],
        ["Reproducibility", "Baseline stored in Supabase"],
    ],
    col_widths=[2.4, 4.2])
spacer(doc, 6)
h2(doc, "Automated smoke testing")
para(doc, "A 13-endpoint live smoke suite runs against production and currently passes 13/13, covering core pages, stock detail (US + India), screener, portfolio, trading dashboard, news, methodology, and pricing.", size=10.5, space_after=6)
h2(doc, "Intellectual property & patent moat")
para(doc, "The platform benefits from a developing IP moat. A provisional patent application has been filed with the Indian Patent Office, covering three core inventive elements:", size=10.5, space_after=4)
bullet(doc, "The PARA-DEBATE adversarial multi-agent debate architecture (a structured “Devil's Advocate” agent that steelmans the bear case before synthesis).")
bullet(doc, "The HMM-based regime-adaptive factor-weighting engine (factor weights shift dynamically across Risk-On / Risk-Off / Bear / Late-Cycle market states).")
bullet(doc, "The NSE delivery_pct India-specific liquidity conviction signal — a market-microstructure factor not present in any comparable system.")
para(doc, "This establishes patent-pending status in India, with a PCT filing contemplated to extend protection across 150+ jurisdictions. In parallel, the PARA-DEBATE™ trademark is being registered in India under Class 42. Backend trade secrets — per-regime factor weights, HMM transition probabilities, agent-prompt architecture, and the IRS% composite formula — sit alongside the patent as a second, non-public layer of protection.", size=10.5, space_after=6)
h2(doc, "Regulatory posture")
para(doc, "Methodology and backtest results are published on a public methodology page in a legally-cautious manner (no guaranteed-return claims). SEBI-compliance considerations for India have been researched.", size=10.5, color=SLATE, space_after=6)

# ===== 10. ENGINEERING INVESTMENT =====
h1(doc, 10, "Engineering Investment (Cost Breakdown)")
para(doc, "The asking price is grounded in what it would cost an acquirer to recreate an equivalent system from scratch — in money, time, and risk. The figures below are rebuild-cost estimates based on senior-engineer effort per subsystem, in person-months. They describe scope, not proprietary internals.", space_after=8)
h2(doc, "Subsystem rebuild effort")
styled_table(doc,
    ["Subsystem", "Scope (no IP detail)", "Senior eng-months"],
    [
        ["Backend API", "33 async routers, auth, quota, rate limit, sessions", "4–6"],
        ["Multi-agent research engine", "Adversarial committee, synthesis, verdict", "3–4"],
        ["Written research analyst", "Query routing, data validation, live price", "2–3"],
        ["Quantitative scoring engine", "Multi-factor + IRS% + regime + walk-forward", "2–3"],
        ["Data pipeline", "6-source ingestion, US+India parity, caching", "3–4"],
        ["Voice agent #1 (PM)", "LiveKit, ~20 tools, STT/TTS, upload, whiteboard", "2–3"],
        ["Voice agent #2 (companion)", "Ambient, wake word, page context, briefing", "1–2"],
        ["Frontend web app", "38 pages, PWA, voice UI, charts, auth flows", "3–4"],
        ["Live trading dashboard", "Real-time matrix, SSE, equity curve, tape", "1–2"],
        ["DevOps & infra", "7-service deploy, CI/CD, 4 cron, secrets, monitoring", "1–2"],
        ["Security hardening", "RLS, IDOR, CSP, audit log, fuses, dep-audit, IR", "1–2"],
        ["Testing", "116+ tests, live smoke, backtest harness", "1–2"],
        ["TOTAL", "24–37 senior eng-months", ""],
    ],
    col_widths=[2.0, 3.4, 1.2])
spacer(doc, 8)
h2(doc, "Translated to cost")
para(doc, "At senior-engineer rates:", size=10.5, space_after=4)
bullet(doc, " (₹2–3L/month): ₹48L–1.1Cr in pure labor.", bold_lead="India senior")
bullet(doc, " ($12–18k/month): $290k–$665k  (₹2.4–5.6Cr).", bold_lead="US senior")
spacer(doc, 4)
para(doc, "These are labor-only figures. They exclude infrastructure spend to date, domain and brand assets, patent filing fees and legal, 6–12 months of calendar opportunity cost, and build risk (a from-scratch rebuild has material probability of schedule overrun, data-source breakage, and integration failure that a live system has already absorbed).", size=10.5, color=SLATE, space_after=8)
h2(doc, "The “live and de-risked” premium")
para(doc, "An acquirer purchasing a running system avoids 6–12 months of build calendar, data-source integration risk (the platform already solved multiple provider failures, rate limits, and US/India market divergences), voice-stack integration risk (LiveKit + STT + LLM + TTS pipelines are notoriously flaky to stabilize), and security debt accumulation. Conservatively, the de-risked + live + calendar-saved premium is 1.5–3x raw rebuild cost.", size=10.5, space_after=8)

# ===== 11. VALUATION RATIONALE =====
h1(doc, 11, "Valuation Rationale")
para(doc, "The single firm ask of ₹3 Crore sits at the lower end of the strategic / IP value band and is supported by every approach below.", space_after=8)
styled_table(doc,
    ["Approach", "Result"],
    [
        ["Raw rebuild cost (India senior)", "₹48L–1.1Cr"],
        ["Raw rebuild cost (US senior)", "₹2.4–5.6Cr"],
        ["Rebuild + live/de-risked premium (1.5–3x)", "₹1.5–3.4Cr"],
        ["Strategic / IP value (IRS% + patent + voice moat)", "₹3–5Cr"],
        ["Revenue multiple", "Not applicable (pre-revenue; see Revenue Model Projection)"],
    ],
    col_widths=[3.6, 3.0])
spacer(doc, 6)
h2(doc, "Pre-revenue by design — a clean greenfield, not a gap")
para(doc, "The platform has been deliberately kept pre-revenue, with the seller prioritising technical robustness, IP defensibility, and backtest validation over early-stage commercialisation. Payments infrastructure, including Stripe integration, is already in place and activation-ready, enabling an acquirer to capture 100% of monetisation upside from closing. With no legacy pricing obligations, no embedded churn, and no inherited customer-support burden, the asset offers a clean, greenfield revenue build on top of a production-ready, patent-pending platform.", size=10.5, space_after=6)
para(doc, "Accordingly, valuation is anchored on asset/rebuild + IP premium, not on a revenue multiple. A revenue multiple would apply only after the acquirer monetizes — see the illustrative ROI model in the next section.", size=10.5, color=SLATE, space_after=6)

# ===== 12. REVENUE MODEL PROJECTION =====
h1(doc, 12, "Revenue Model Projection (Illustrative)")
para(doc, "The platform is pre-revenue today, but the monetisation engine is built and activation-ready. The scenarios below are illustrative, not guaranteed — they exist to give an acquirer a mental model for return on investment.", space_after=8)
styled_table(doc,
    ["Scenario", "Assumption", "Indicative ARR"],
    [
        ["Conservative", "2,000 paying subscribers @ $9.99/mo", "~$240K ARR"],
        ["Base", "5,000 subscribers + B2B API tier", "~$800K ARR"],
        ["Upside", "50,000 subscribers via broker partnership", "~$6M ARR"],
    ],
    col_widths=[1.5, 3.6, 1.5])
spacer(doc, 6)
para(doc, "For a strategic buyer with an existing distribution channel — a broker or wealth platform — the upside scenario makes payback effectively immediate. For a standalone buyer, the base-case ARR recovers the purchase price well inside the first year of monetisation (illustrative, gross of operating costs).", size=10.5, space_after=4)
para(doc, "These projections are illustrative scenarios provided for buyer modelling only. They are not forecasts, guarantees, or representations of future performance.", size=9, italic=True, color=SLATE, space_after=6)

# ===== 13. ASKING PRICE =====
h1(doc, 13, "Asking Price")
pp = doc.add_paragraph(); pp.alignment = CENTER
pp.paragraph_format.space_before = Pt(6); pp.paragraph_format.space_after = Pt(6)
shade_para(pp, NAVY_HEX); para_box_border(pp, color=GOLD_HEX, sz="12")
re = pp.add_run("₹3 Crore"); style_run(re, 30, True, GOLD)
re2 = pp.add_run("    ·    a single, firm ask"); style_run(re2, 12, False, WHITE)
spacer(doc, 6)
para(doc, "The price is anchored on rebuild cost plus the live / de-risked premium plus strategic IP value (see Engineering Investment and Valuation Rationale), and sits at the lower end of the strategic-value band. It is a firm ask for a serious strategic acquirer — a brokerage, fund, or fintech — not an opening number in a negotiation ladder.", size=10.5, color=SLATE, space_after=6)

# ===== 14. WHAT TRANSFERS =====
h1(doc, 14, "What Transfers on Acquisition")
bullet(doc, "Domain (neuralquant.co) and brand assets")
bullet(doc, "Full source code (monorepo) under NDA / LOI")
bullet(doc, "All infrastructure accounts (Render, Vercel, Supabase, Railway, LiveKit) — credentials transferred securely")
bullet(doc, "Data pipeline configurations and cached market data")
bullet(doc, "Patent application (assignment to acquirer)")
bullet(doc, "Methodology and backtest baselines")
bullet(doc, "Operations runbooks and security documentation")
bullet(doc, "A knowledge-transfer period (suggested 2–4 weeks) to hand over operations")

# ===== 15. EXCLUSIONS =====
h1(doc, 15, "Exclusions / Not Included")
bullet(doc, "Third-party API keys are not transferred; the acquirer provisions its own accounts (FMP, Finnhub, ElevenLabs, Deepgram, Anthropic, AWS, LiveKit, OpenBB). This is standard and avoids key-abuse liability.")
bullet(doc, "No customer / user data is represented as included (the platform is pre-revenue with minimal user base).")
bullet(doc, "No revenue or traffic representations are made.")
bullet(doc, "Proprietary algorithm internals, agent prompts, and scoring formulas are disclosed only during due diligence under NDA.")

# ===== 16. DD CHECKLIST =====
h1(doc, 16, "Buyer's Due-Diligence Checklist")
checkbox(doc, "Live site walk-through (neuralquant.co) — all flagship flows demonstrated live")
checkbox(doc, "Smoke test suite execution (13/13 passing)")
checkbox(doc, "Backtest reproduction (baseline stored in Supabase)")
checkbox(doc, "Patent filing receipt and invention disclosure (under NDA)")
checkbox(doc, "Service-by-service infra review (under NDA)")
checkbox(doc, "Security audit documentation (RLS, IDOR, CSP, audit log)")
checkbox(doc, "Code structure review (under NDA)")
checkbox(doc, "Methodology page (public, no NDA needed)")

# ===== 17. WHY FAIR =====
h1(doc, 17, "Why The Price Is Fair")
bullet(doc, "You are buying time. 6–12 months of build calendar and the data/voice integration scars that go with it.", bold_lead="1. ")
bullet(doc, "You are buying working IP. A unique scoring methodology and an adversarial multi-agent design that no competitor currently ships, plus a pending patent.", bold_lead="2. ")
bullet(doc, "You are buying a de-risked live system, not a bet on whether it can be built.", bold_lead="3. ")
bullet(doc, "The ask is below US rebuild cost plus the live/IP premium. Even at ₹3 Cr, the acquirer pays less than recreating it would cost in senior labor plus the de-risked premium.", bold_lead="4. ")
bullet(doc, "Strategic upside is large. For a brokerage, fund, or fintech, the voice + IRS% + adversarial-trust combination is a differentiated feature that would take a competitor a full cycle to copy — and the patent raises that cost further.", bold_lead="5. ")

# ===== 18. NEXT STEPS =====
h1(doc, 18, "Next Steps")
bullet(doc, "Sign a mutual NDA.")
bullet(doc, "Live product demonstration (remote, ~45 minutes).")
bullet(doc, "Letter of intent → full source and infra access for due diligence.")
bullet(doc, "Knowledge-transfer period post-close.")
spacer(doc, 8)
para(doc, "Contact: Satyam Das  —  satyamdas03@gmail.com",
     size=12, bold=True, color=NAVY, align=CENTER, space_before=10, space_after=10)

rule2 = doc.add_paragraph(); rule2.alignment = CENTER
pPr = rule2._p.get_or_add_pPr()
pbdr = OxmlElement('w:pBdr')
bottom = OxmlElement('w:bottom')
bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '12')
bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), CYAN_HEX)
pbdr.append(bottom); pPr.append(pbdr)
rr2 = rule2.add_run("    "); style_run(rr2, 2, False, WHITE)

para(doc, "This document is a sales prospectus, not a binding offer. Final terms are set in a definitive acquisition agreement. Valuations are the seller's estimate and have not been independently appraised; acquirers are encouraged to form their own view during due diligence.",
     size=9, italic=True, color=SLATE, align=CENTER, space_before=8, space_after=0)

# ---------- save ----------
out = r"C:\Users\point\projects\stockpredictor\docs\SALE_PROSPECTUS.docx"
doc.save(out)
print("SAVED:", out)
