"""Generate a styled .docx StockEdge partnership prospectus for NeuralQuant.
Output: docs/SALES PROSPECTUS/STOCKEDGE_PARTNERSHIP_PROSPECTUS.docx
Brand palette: deep navy + electric cyan accent + warm gold.
Tables-first, glance-readable. Lead offer: 12-month licensing pilot + call option.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------- brand palette ----------
NAVY = RGBColor(0x0A, 0x1B, 0x2E)
CYAN = RGBColor(0x00, 0xB4, 0xD8)
GOLD = RGBColor(0xC9, 0xA2, 0x27)
SLATE = RGBColor(0x3A, 0x4A, 0x5E)
LIGHT = RGBColor(0xF2, 0xF5, 0xF8)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK = RGBColor(0x12, 0x1A, 0x24)

NAVY_HEX = "0A1B2E"
CYAN_HEX = "00B4D8"
LIGHT_HEX = "EEF3F7"
GOLD_HEX = "C9A227"

CENTER = WD_ALIGN_PARAGRAPH.CENTER

# ---------- helpers ----------
def set_cell_shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, color="C7D2DC", sz="4"):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'single')
        e.set(qn('w:sz'), sz)
        e.set(qn('w:space'), '0')
        e.set(qn('w:color'), color)
        borders.append(e)
    tcPr.append(borders)

def style_run(run, size=11, bold=False, color=DARK, font="Calibri", italic=False):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    run.font.name = font

def shade_para(p, hex_fill):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_fill)
    pPr.append(shd)

def para_box_border(p, color=CYAN_HEX, sz="6"):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    for edge in ('top', 'left', 'bottom', 'right'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), sz)
        e.set(qn('w:space'), '6'); e.set(qn('w:color'), color)
        pbdr.append(e)
    pPr.append(pbdr)

def add_field(paragraph, field_code):
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve'); instrText.text = field_code
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar2)
    return run

def para(doc, text, size=11, bold=False, color=DARK, align=None, space_after=6, space_before=0, italic=False, font="Calibri"):
    p = doc.add_paragraph()
    if align is not None: p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    r = p.add_run(text); style_run(r, size, bold, color, font, italic)
    return p

def h1(doc, num, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(f"{num}.  {text}")
    style_run(r, 16, True, NAVY, "Calibri")
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '4'); bottom.set(qn('w:color'), CYAN_HEX)
    pbdr.append(bottom); pPr.append(pbdr)
    return p

def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    style_run(r, 12.5, True, SLATE, "Calibri")
    return p

def h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    style_run(r, 11.5, True, SLATE, "Calibri")
    return p

def bullet(doc, text, bold_lead=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    if bold_lead:
        r1 = p.add_run(bold_lead); style_run(r1, 10.5, True, DARK)
        r2 = p.add_run(text); style_run(r2, 10.5, False, DARK)
    else:
        r = p.add_run(text); style_run(r, 10.5, False, DARK)
    return p

def checkbox(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.3)
    r0 = p.add_run("☐  "); style_run(r0, 11, False, CYAN, "Calibri")
    r = p.add_run(text); style_run(r, 10.5, False, DARK)
    return p

def styled_table(doc, headers, rows, col_widths=None, header_fill=NAVY_HEX, zebra=True):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for i, htext in enumerate(headers):
        cell = t.rows[0].cells[i]
        set_cell_shade(cell, header_fill)
        set_cell_borders(cell, color="0A1B2E", sz="4")
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = cell.paragraphs[0].add_run(htext)
        style_run(r, 10.5, True, WHITE, "Calibri")
        cell.paragraphs[0].paragraph_format.space_after = Pt(2)
        cell.paragraphs[0].paragraph_format.space_before = Pt(2)
    for ridx, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            c = cells[i]
            if zebra and ridx % 2 == 1:
                set_cell_shade(c, LIGHT_HEX)
            set_cell_borders(c)
            c.paragraphs[0].paragraph_format.space_after = Pt(2)
            c.paragraphs[0].paragraph_format.space_before = Pt(2)
            r = c.paragraphs[0].add_run(str(val))
            style_run(r, 10, False, DARK, "Calibri")
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    return t

def label_table(doc, rows, label_w=1.7, body_w=4.9, header_fill=NAVY_HEX):
    t = doc.add_table(rows=0, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for label, body in rows:
        cells = t.add_row().cells
        c0, c1 = cells[0], cells[1]
        set_cell_shade(c0, header_fill); set_cell_borders(c0, color="0A1B2E", sz="4")
        set_cell_borders(c1)
        c0.paragraphs[0].paragraph_format.space_after = Pt(2); c0.paragraphs[0].paragraph_format.space_before = Pt(2)
        c1.paragraphs[0].paragraph_format.space_after = Pt(2); c1.paragraphs[0].paragraph_format.space_before = Pt(2)
        r0 = c0.paragraphs[0].add_run(label); style_run(r0, 10, True, WHITE)
        r1 = c1.paragraphs[0].add_run(body); style_run(r1, 10, False, DARK)
    for row in t.rows:
        row.cells[0].width = Inches(label_w); row.cells[1].width = Inches(body_w)
    return t

def spacer(doc, pts=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(pts)
    p.paragraph_format.space_before = Pt(0)
    return p

def code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.2)
    shade_para(p, LIGHT_HEX)
    para_box_border(p, color=CYAN_HEX, sz="4")
    r = p.add_run(text)
    style_run(r, 9, False, SLATE, "Consolas")
    return p

# ---------- build document ----------
doc = Document()

for s in doc.sections:
    s.top_margin = Inches(0.8)
    s.bottom_margin = Inches(0.8)
    s.left_margin = Inches(0.9)
    s.right_margin = Inches(0.9)

normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)
normal.font.color.rgb = DARK

# ===== COVER PAGE =====
band = doc.add_paragraph()
band.paragraph_format.space_after = Pt(0)
pPr = band._p.get_or_add_pPr()
shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), NAVY_HEX); pPr.append(shd)
rb = band.add_run("  "); style_run(rb, 6, False, NAVY)

para(doc, "CONFIDENTIAL  ·  STRATEGIC PARTNERSHIP PROPOSAL", size=10, bold=True, color=CYAN,
     align=CENTER, space_before=100, space_after=4)

title = doc.add_paragraph(); title.alignment = CENTER
title.paragraph_format.space_after = Pt(2)
r = title.add_run("NeuralQuant + StockEdge"); style_run(r, 40, True, NAVY, "Calibri")

sub = doc.add_paragraph(); sub.alignment = CENTER
sub.paragraph_format.space_after = Pt(2)
r = sub.add_run("AI Research Pilot + Call Option"); style_run(r, 18, False, SLATE, "Calibri")

rule = doc.add_paragraph(); rule.alignment = CENTER
rule.paragraph_format.space_after = Pt(14); rule.paragraph_format.space_before = Pt(8)
pPr = rule._p.get_or_add_pPr()
pbdr = OxmlElement('w:pBdr')
bottom = OxmlElement('w:bottom')
bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '18')
bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), GOLD_HEX)
pbdr.append(bottom); pPr.append(pbdr)
rr = rule.add_run("    "); style_run(rr, 2, False, WHITE)

para(doc, "A 12-month exclusive licensing pilot to embed PARA-DEBATE™ + IRS% inside StockEdge",
     size=11, bold=False, color=SLATE, align=CENTER, space_after=120)

pp = doc.add_paragraph(); pp.alignment = CENTER
pp.paragraph_format.space_before = Pt(6); pp.paragraph_format.space_after = Pt(6)
shade_para(pp, NAVY_HEX); para_box_border(pp, color=GOLD_HEX, sz="10")
re = pp.add_run("Lead offer: ₹50 lakh upfront + 15% revenue share + ₹2.5 Cr call option"); style_run(re, 15, True, GOLD)

meta = doc.add_table(rows=5, cols=2)
meta.alignment = WD_TABLE_ALIGNMENT.CENTER
meta_data = [
    ("Prepared", "June 2026"),
    ("From", "Satyam Das  —  satyamdas03@gmail.com"),
    ("To", "StockEdge Fintech Pvt Ltd"),
    ("Audience", "Vineet Patawari, CEO  ·  Vivek Bajaj, MD"),
    ("Status", "Live platform  ·  v4.1.0  ·  13/13 smoke tests passing"),
]
for i, (k, v) in enumerate(meta_data):
    c0 = meta.rows[i].cells[0]; c1 = meta.rows[i].cells[1]
    c0.width = Inches(1.4); c1.width = Inches(4.2)
    set_cell_borders(c0, color="FFFFFF", sz="2"); set_cell_borders(c1, color="FFFFFF", sz="2")
    r0 = c0.paragraphs[0].add_run(k.upper()); style_run(r0, 9, True, CYAN)
    r1 = c1.paragraphs[0].add_run(v); style_run(r1, 10.5, False, DARK)
    c0.paragraphs[0].paragraph_format.space_after = Pt(4)
    c1.paragraphs[0].paragraph_format.space_after = Pt(4)

para(doc, "Strictly confidential — transmit only under signed mutual NDA.",
     size=9, italic=True, color=SLATE, align=CENTER, space_before=120, space_after=0)

body_section = doc.sections[0]
footer = body_section.footer
fp = footer.paragraphs[0]
fp.alignment = CENTER
r = fp.add_run("NeuralQuant + StockEdge — Strategic Partnership Proposal   |   Confidential   |   Page ")
style_run(r, 8, False, SLATE)
add_field(fp, "PAGE")
r2 = fp.add_run(" of "); style_run(r2, 8, False, SLATE)
add_field(fp, "NUMPAGES")

doc.add_page_break()

# ===== 1. EXECUTIVE SUMMARY =====
h1(doc, 1, "Executive Summary")
para(doc, "NeuralQuant is a production-grade, multi-agent AI stock-research platform covering the US and India equity markets (~949 stocks). It is live and operational today at neuralquant.co — not a prototype, not a slide deck.", space_after=6)
para(doc, "StockEdge has built India's largest retail research distribution engine: 4M+ registered users, SEBI-registered Research Analyst and Investment Adviser licenses, and a publicly stated ambition to create a 'Bloomberg-for-Bharat' ecosystem. NeuralQuant has built the reasoning layer that turns data into defended, risk-aware investment conclusions.", space_after=6)
para(doc, "This proposal is not a request for a ₹3 Crore cash acquisition on Day 1. It is a low-risk, gated partnership:", space_after=6)
bullet(doc, "12-month exclusive pilot: embed NeuralQuant's PARA-DEBATE™ + IRS% engine inside StockEdge Pro / Club / Investment Cases.")
bullet(doc, "Small upfront license fee to cover integration and exclusivity.")
bullet(doc, "Revenue-share kicker tied only to the incremental subscribers the AI tier converts.")
bullet(doc, "Call option: StockEdge can acquire the IP, codebase, and domain once the pilot proves conversion.")
spacer(doc, 8)
label_table(doc, [
    ("What you get", "A live, patent-pending AI research engine that fills four visible product gaps in StockEdge today."),
    ("What we get", "A low-risk pilot fee + a revenue share on the AI tier we help create + a path to full acquisition."),
    ("What StockEdge users get", "Interactive 'why this stock?' answers, adversarial risk sections, regime-aware IRS% scores, and US+IN dual-market coverage."),
    ("What regulators see", "Research analysis produced under StockEdge's existing SEBI RA/IA licenses — no new registration required."),
])

# ===== 2. STRATEGIC FIT =====
h1(doc, 2, "Why NeuralQuant + StockEdge Makes Strategic Sense")
h2(doc, "2.1 StockEdge's ambition is NeuralQuant's exact output")
para(doc, "Vivek Bajaj has repeatedly said his life goal is to build a Bloomberg-like ecosystem for Indian retail investors. Bloomberg provides data; NeuralQuant provides the reasoning and risk discipline on top of data — bull case, adversarial bear case, regime-aware scoring, and position-sizing guidance. The combination gives StockEdge a genuine 'full stack from data to decision.'", space_after=6)

h2(doc, "2.2 Four gaps in StockEdge today that NeuralQuant fills")
styled_table(doc,
    ["StockEdge gap today", "NeuralQuant capability", "Why it matters"],
    [
        ["No LLM 'Ask AI' research assistant", "Ask Morgan + PARA-DEBATE agents", "Turns static scans into interactive 'why this stock?' answers"],
        ["No adversarial risk section", "PARA-DEBATE™ 6+1 agent committee with mandated BEAR", "Produces the counter-argument section StockEdge lacks"],
        ["No regime-aware position-sizing metric", "IRS% (Investment-Ready Score)", "Complements RS55/RSI with macro-adaptive conviction"],
        ["No US equity coverage", "US + India dual-market pipeline", "Niche differentiator for HNIs and Club subscribers"],
    ],
    col_widths=[1.7, 2.1, 2.8])

h2(doc, "2.3 Revenue synergy paths")
styled_table(doc,
    ["StockEdge revenue line", "How NeuralQuant strengthens it"],
    [
        ["StockEdge Pro / Premium (~₹11,989/yr)", "Bundle PARA-DEBATE reports for NSE 200 as a higher-tier 'Pro AI' upsell"],
        ["Investment Cases / smallcase (₹3,000–12,000/yr)", "Use IRS% + regime output to justify rebalancing in model portfolios"],
        ["StockEdge Club (₹23,989/yr)", "Live 'Ask the AI Analyst' sessions on audience-requested stocks"],
        ["Kotak Neo / B2B distribution", "Resell the engine to broker partners, mirroring the 2021 Kotak partnership"],
    ],
    col_widths=[2.3, 4.3])

h2(doc, "2.4 Regulatory complementarity")
para(doc, "StockEdge holds both SEBI registrations that NeuralQuant lacks:", space_after=4)
bullet(doc, "Research Analyst (RA): INH300007493")
bullet(doc, "Investment Adviser (IA): INA000017781")
para(doc, "By embedding NeuralQuant as a research-analysis tool under StockEdge's existing licenses, the output is immediately SEBI-compliant — with StockEdge's compliance officer sign-off, disclaimers, and no guaranteed-return language. NeuralQuant does not need its own RA/IA registration.", size=10.5, color=SLATE, space_after=6)

# ===== 3. WHAT NEURALQUANT IS =====
h1(doc, 3, "What NeuralQuant Is")
h2(doc, "3.1 At a glance")
styled_table(doc,
    ["Dimension", "Scale"],
    [
        ["Live site", "neuralquant.co (Vercel, auto-deployed)"],
        ["API", "neuralquant.onrender.com (FastAPI, v4.1.0)"],
        ["Backend routers", "33"],
        ["Web pages", "38"],
        ["Running services", "7 (4 Render, 1 Vercel, 1 Railway, 1 Supabase)"],
        ["Markets", "US + India"],
        ["Stock universe", "~949 stocks (502 India + 447 US)"],
        ["AI agents", "8 distinct agents"],
        ["Live voice interfaces", "2 (voice PM + ambient companion)"],
        ["Production tests", "148 backend tests passing, 13/13 live smoke suite"],
        ["Patent status", "Provisional application filed (India), PCT contemplated"],
    ],
    col_widths=[2.0, 4.6])

h2(doc, "3.2 Flagship capabilities")
styled_table(doc,
    ["Capability", "What it does"],
    [
        ["PARA-DEBATE™ adversarial committee ★", "6 specialist agents + mandated BEAR + Head Analyst synthesis. Every conclusion stress-tested by a steelmanned bear case. Prevents consensus herding and hallucination."],
        ["IRS% scoring engine ★", "Five-factor composite (quality, momentum, value, low-vol, regime) + NSE Bhavcopy delivery_pct for India — unique market-microstructure signal."],
        ["HMM regime detection ★", "4-state Hidden Markov Model dynamically reweighting factor exposure by market regime."],
        ["Ask Morgan", "Written AI analyst with live price injection, clarification questions, and numeric reconciliation against a [VERIFIED] data layer."],
        ["Veronica + QuantAstra", "Two live voice agents via LiveKit, Deepgram STT, ElevenLabs TTS, and Anthropic/Bedrock LLMs."],
        ["Hermes live trading dashboard", "Real-time paper-trading matrix with equity curve, trade tape, and strategy-reflection SSE stream."],
    ],
    col_widths=[2.1, 4.5])

h2(doc, "3.3 Validation & track record")
styled_table(doc,
    ["Metric", "Result"],
    [
        ["Q1 FY27 benchmark", "NIFTY50: −6.38%"],
        ["Alpha vs NIFTY50", "+12.69% to +14.83%"],
        ["Hit rate", "87–91%"],
        ["Live smoke suite", "13/13 passing"],
        ["Backtest reproducibility", "Baseline stored in Supabase"],
    ],
    col_widths=[2.4, 4.2])

# ===== 4. DEAL STRUCTURE =====
h1(doc, 4, "The Proposed Deal Structure")
h2(doc, "4.1 Primary recommendation: 12-month licensing pilot + call option")
styled_table(doc,
    ["Component", "Term"],
    [
        ["Upfront license fee", "₹50 lakh for exclusive India-market integration rights to PARA-DEBATE™ + IRS%"],
        ["Universe covered", "NSE 200 + BSE 500 (expandable by mutual agreement)"],
        ["Revenue share", "15% of incremental ARR from any new AI tier launched using the engine"],
        ["Revenue-share cap", "₹2 Crore over 24 months"],
        ["Pilot gates", "10,000 MAU; 1,000 paid upgrades; avg report latency < 30s"],
        ["Call option", "Acquire NeuralQuant IP, codebase, and neuralquant.co for ₹2.5 Crore within 24 months"],
        ["License fee credit", "₹50 lakh upfront credited against call-option purchase price"],
        ["SEBI wrapper", "All outputs labeled under StockEdge RA INH300007493 with required disclaimers"],
    ],
    col_widths=[2.4, 4.2])

h2(doc, "4.2 Why this structure fits StockEdge")
bullet(doc, "Low upfront cash: ₹50 lakh is less than two senior engineers for a year — minimal balance-sheet impact.")
bullet(doc, "Aligned incentives: most of NeuralQuant's compensation is a revenue share on the new AI tier.")
bullet(doc, "Risk reversal: StockEdge only exercises the call option if pilot data justifies it.")
bullet(doc, "Regulatory safety: engine sits under StockEdge's existing RA/IA licenses.")
bullet(doc, "Exclusivity: prevents NeuralQuant from licensing the same engine to a competing Indian platform during the pilot.")

h2(doc, "4.3 Fallback structures")
styled_table(doc,
    ["Structure", "Mechanics", "When it suits StockEdge"],
    [
        ["Strategic investment + distribution", "₹5–8 Cr for 15–20% stake; exclusive integration; board observer", "Deeper alignment without full acquisition risk"],
        ["Acqui-hire / IP buyout", "₹1.5–2 Cr cash for team + codebase; founder stays 6–12 months", "Absorb team and fold IP in-house quickly"],
        ["Full acquisition with earnout", "₹1 Cr upfront + ₹2 Cr tied to AI-tier conversions over 18 months", "Confident in immediate product-market fit"],
        ["Stock swap", "NeuralQuant equity swapped into KIPL shares", "Only if near-term IPO path is credible"],
    ],
    col_widths=[1.8, 2.8, 2.0])

# ===== 5. SEBI INTEGRATION SPEC =====
h1(doc, 5, "SEBI-Compliant Integration Specification")
h2(doc, "5.1 Regulatory premise")
para(doc, "StockEdge is the licensed entity. NeuralQuant is the technology provider. The AI engine produces research analysis, not personalized investment advice. Final labeling, disclaimers, and any 'buy/sell/hold' language are controlled by StockEdge under its RA/IA registrations.", space_after=6)

h2(doc, "5.2 Output labeling (every user-facing report)")
code_block(doc, "Research analysis powered by NeuralQuant.\nReviewed under StockEdge Research Analyst registration INH300007493.\nThis is not investment advice. Past performance does not guarantee future results.\nConsult a SEBI-registered investment adviser before acting.")

h2(doc, "5.3 Compliance controls")
styled_table(doc,
    ["Control", "Implementation"],
    [
        ["No guaranteed-return claims", "Hard-coded disclaimer on every AI-generated output"],
        ["Principal-officer sign-off", "StockEdge compliance officer approves AI-tier output template before launch"],
        ["Record-keeping", "All AI reports logged with timestamp, ticker, input signals, and version hash"],
        ["Research vs advisory segregation", "AI tier labeled 'Research' only; personalized advisory remains with StockEdge IA team"],
        ["MITC disclosures", "Included in AI-tier subscription terms"],
    ],
    col_widths=[2.2, 4.4])

# ===== 6. INTEGRATION PLAN =====
h1(doc, 6, "12-Month Integration Plan")
styled_table(doc,
    ["Phase", "Timeline", "Activities"],
    [
        ["Phase 1: Embed", "Months 1–2", "Sandbox integration; private /partners/stockedge endpoint; compliance review; internal soft launch"],
        ["Phase 2: Pilot Tier Launch", "Months 3–6", "Launch 'StockEdge Pro AI' or 'Investment Cases AI Rationale'; top 50 NSE stocks → NSE 200; track MAU, upgrades, latency"],
        ["Phase 3: Scale Decision", "Months 7–12", "Expand to BSE 500 / Club live sessions / US-stock HNI tier; StockEdge decides on call option"],
        ["Handover", "Post-LOI", "2–4 weeks knowledge transfer; architecture, operations, and security documentation"],
    ],
    col_widths=[1.6, 1.3, 3.7])

# ===== 7. OUTBOUND MATERIALS =====
h1(doc, 7, "Outbound Materials")
h2(doc, "7.1 LinkedIn message to Vineet Patawari (CEO)")
code_block(doc, "Hi Vineet — long-time admirer of what you and Vivek have built at StockEdge.\n\nWe have built an adversarial AI research committee (PARA-DEBATE) that forces a bull-vs-bear debate before any stock conclusion — exactly the 'noise removal' StockEdge stands for. It is live with dual-market US+IN coverage, an India-specific IRS% score, and a Q1FY27 backtest showing +12.69% to +14.83% alpha vs NIFTY50.\n\nI am not looking for a cash-heavy acquisition day one. I would love to explore a 12-month pilot inside StockEdge Pro that converts your 4M users into AI-tier subscribers, with an option to buy the IP later.\n\nWorth a 15-minute call?")

h2(doc, "7.2 LinkedIn message to Vivek Bajaj (MD)")
code_block(doc, "Hi Vivek — your Bloomberg-for-Bharat vision is why I started NeuralQuant.\n\nWe built the reasoning layer that turns market data into defended conclusions: a quantified adversarial committee, regime-adaptive scoring, and an India-specific IRS% metric. Live at neuralquant.co, 949 tickers, 13/13 smoke tests passing.\n\nI would welcome the chance to show you how it could sit inside StockEdge's existing SEBI RA/IA framework as a 12-month pilot — low risk, aligned incentives, and only an acquisition if your users actually pay for it.\n\nMay I send a one-page memo?")

h2(doc, "7.3 Email subject lines")
bullet(doc, "12-month AI research pilot for StockEdge Pro — low upfront, aligned upside")
bullet(doc, "NeuralQuant + StockEdge: a Bloomberg-for-Bharat reasoning layer")
bullet(doc, "Adversarial AI research engine — 12.69–14.83% alpha vs NIFTY50 — pilot proposal")

# ===== 8. DEMO AGENDA =====
h1(doc, 8, "10-Minute Demo Agenda")
styled_table(doc,
    ["Time", "What to show"],
    [
        ["0:00–0:30", "Live site, 949 tickers, 13/13 smoke tests, v4.1.0"],
        ["0:30–2:30", "Ask Morgan: natural-language question on RELIANCE / TCS / HDFCBANK"],
        ["2:30–5:30", "PARA-DEBATE on the same stock — 6 agents, mandated bear, consensus, risk section"],
        ["5:30–7:00", "IRS% / screener — rank NSE 200, show regime detection shifting factor weights"],
        ["7:00–8:30", "Backtest evidence: Q1FY27 alpha +12.69–14.83%, hit rate 87–91%"],
        ["8:30–10:00", "Proposed pilot: integration, SEBI umbrella, ₹50L + 15% + call option"],
    ],
    col_widths=[1.2, 5.4])

# ===== 9. NEGOTIATION =====
h1(doc, 9, "Negotiation Talking Points")
h2(doc, "9.1 Five sentences that will resonate")
bullet(doc, "You have built the largest retail research distribution engine in India; we have built the reasoning layer that turns your scans into 'why' — together you own the full stack from data to decision.")
bullet(doc, "PARA-DEBATE is not a chatbot; it is a quantified adversarial committee that forces the bear case before any buy idea reaches your user — that protects your brand from bad calls.")
bullet(doc, "Your RS55 and RSI already teach discipline; IRS% adds regime-adaptive position sizing so your users know when the macro wind is against them.")
bullet(doc, "You do not need to spend 24–37 senior engineer-months and ₹1.1 Cr+ in labor to recreate this; it is live today and can be embedded under your existing SEBI licenses.")
bullet(doc, "We are not asking you to bet the company on a ₹3 Cr acquisition — we are asking for a 12-month pilot, and you only buy the IP if your users actually pay for it.")

h2(doc, "9.2 Objections and answers")
styled_table(doc,
    ["Objection", "Answer"],
    [
        ["Why can't we build this ourselves?", "A reliable adversarial multi-agent system with hallucination guards, numeric reconciliation, and a live dual-market data pipeline has taken 80+ iterations and 126 documented bug cycles. Rebuild cost: 24–37 senior eng-months plus integration scars. License first, validate, then decide."],
        ["Is this SEBI-compliant?", "NeuralQuant does not hold RA/IA licenses; StockEdge does. We embed the engine as a research-analysis tool under your existing registrations, with your compliance officer's sign-off and required disclaimers."],
        ["You have no revenue. Why pay anything?", "You are paying for time-to-market and IP protection. ₹50 lakh is less than two senior engineers for a year; most compensation is a 15% revenue share on the new AI tier."],
        ["Will this dilute my brand as the expert?", "The AI augments your momentum/RS framework, not replaces it. PARA-DEBATE produces the counter-argument and risk sections your analysts already believe in — faster and at scale."],
        ["What if the pilot doesn't convert?", "Your downside is capped at ₹50 lakh plus integration cost. You keep the learnings. If gates are not met, exclusivity terminates and you walk away with no acquisition obligation."],
    ],
    col_widths=[2.3, 4.3])

# ===== 10. REBUILD COST =====
h1(doc, 10, "Engineering Investment Reference")
para(doc, "The ₹50 lakh pilot fee is a fraction of the cost to recreate even one subsystem.", space_after=6)
styled_table(doc,
    ["Subsystem", "Senior eng-months"],
    [
        ["Backend API (33 routers, auth, quotas, sessions)", "4–6"],
        ["PARA-DEBATE multi-agent engine", "3–4"],
        ["Ask Morgan written analyst", "2–3"],
        ["Quant scoring + IRS% + HMM regime", "2–3"],
        ["Data pipeline (6 sources, US+IN parity)", "3–4"],
        ["Voice PM + ambient companion", "3–5"],
        ["Next.js web app (38 pages, PWA)", "3–4"],
        ["Live trading dashboard", "1–2"],
        ["DevOps, security, testing", "3–5"],
        ["TOTAL", "24–37"],
    ],
    col_widths=[4.8, 1.8])
para(doc, "India senior rates (₹2–3L/mo): ₹48L–1.1Cr. US senior rates ($12–18k/mo): $290k–$665k (₹2.4–5.6Cr). Excludes calendar time, integration scars, data-source failures, voice-stack stabilization, and security hardening.", size=10.5, color=SLATE, space_after=6)

# ===== 11. TRANSFERS / EXCLUSIONS =====
h1(doc, 11, "What Transfers / Exclusions")
h2(doc, "11.1 Transfers (on exercise of call option or IP buyout)")
bullet(doc, "Domain (neuralquant.co) and brand assets")
bullet(doc, "Full source code (monorepo) under NDA / LOI")
bullet(doc, "Infrastructure accounts — credentials transferred securely")
bullet(doc, "Data pipeline configurations and cached market data")
bullet(doc, "Patent application — assignment to StockEdge")
bullet(doc, "Methodology and backtest baselines")
bullet(doc, "Operations and security documentation")
bullet(doc, "Knowledge-transfer period (2–4 weeks)")

h2(doc, "11.2 Exclusions")
bullet(doc, "Third-party API keys (StockEdge provisions its own accounts)")
bullet(doc, "Customer/user data (platform is pre-revenue)")
bullet(doc, "Revenue or traffic representations")
bullet(doc, "Proprietary algorithm internals disclosed only under NDA during due diligence")

# ===== 12. DD CHECKLIST =====
h1(doc, 12, "Due-Diligence Readiness Checklist")
checkbox(doc, "Live product demonstration ready (neuralquant.co)")
checkbox(doc, "13/13 smoke suite passing")
checkbox(doc, "Backtest baseline stored in Supabase (Q1FY27)")
checkbox(doc, "Signed mutual NDA template")
checkbox(doc, "Patent filing receipt + Form 31 grace-period filing status")
checkbox(doc, "PARA-DEBATE™ trademark status confirmation")
checkbox(doc, "One-page 'StockEdge + NeuralQuant' strategic memo")
checkbox(doc, "12-month pilot term sheet draft")
checkbox(doc, "SEBI-compliant output-labeling spec")
checkbox(doc, "Clean repo + CI lint fixed + pending migrations applied")
checkbox(doc, "FMP key rotated + Render env vars documented")

# ===== 13. NEXT STEPS =====
h1(doc, 13, "Next Steps")
bullet(doc, "Sign a mutual NDA — no source code, architecture detail, or patent specifics shared before this.")
bullet(doc, "15-minute product call — demo the live platform on a top NSE stock.")
bullet(doc, "Strategic memo + pilot term sheet — shared for internal StockEdge review.")
bullet(doc, "Letter of Intent — kicks off the 12-month exclusive pilot and integration planning.")
bullet(doc, "Pilot launch — Months 1–2 embed; Months 3–6 pilot tier; Months 7–12 scale decision.")

spacer(doc, 12)
para(doc, "Contact: Satyam Das — satyamdas03@gmail.com",
     size=12, bold=True, color=NAVY, align=CENTER, space_before=10, space_after=10)

rule2 = doc.add_paragraph(); rule2.alignment = CENTER
pPr = rule2._p.get_or_add_pPr()
pbdr = OxmlElement('w:pBdr')
bottom = OxmlElement('w:bottom')
bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '12')
bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), CYAN_HEX)
pbdr.append(bottom); pPr.append(pbdr)
rr2 = rule2.add_run("    "); style_run(rr2, 2, False, WHITE)

para(doc, "This document is a strategic partnership proposal and acquisition prospectus, not a binding offer. Final terms are set in a definitive agreement. Valuations and projections are the seller's estimates and have not been independently appraised; StockEdge is encouraged to form its own view during due diligence.",
     size=9, italic=True, color=SLATE, align=CENTER, space_before=8, space_after=0)

# ---------- save ----------
out = r"C:\Users\point\projects\stockpredictor\docs\SALES PROSPECTUS\STOCKEDGE_PARTNERSHIP_PROSPECTUS.docx"
doc.save(out)
print("SAVED:", out)
